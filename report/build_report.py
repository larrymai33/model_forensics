from __future__ import annotations

import argparse
import re
from pathlib import Path

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "takehome.md"
FIGURES_DIR = ROOT / "figures"
AUDITED_DOCX = ROOT / "odd-number-model-forensics-takehome.docx"
DEFAULT_OUTPUT = ROOT / "generated" / AUDITED_DOCX.name

BLACK = "000000"
MUTED = "555555"
BORDER = "DADCE0"
LIGHT = "F8F9FA"
RED = "A61B1B"
BLUE = "1F5AA6"


def load_build_dependencies() -> None:
    """Load optional document-build packages only after output-policy checks pass."""
    global Image, Document, WD_STYLE_TYPE, WD_CELL_VERTICAL_ALIGNMENT
    global WD_TABLE_ALIGNMENT, WD_ALIGN_PARAGRAPH, OxmlElement, qn
    global Inches, Pt, RGBColor

    try:
        from PIL import Image
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ModuleNotFoundError as exc:
        raise SystemExit(
            "Report build dependencies are missing. "
            "Install the exact pins with: python -m pip install -r report/requirements.txt"
        ) from exc


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 0) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, name="Arial", size=11, bold=None, italic=None, color=BLACK):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, before, after, color in (
        ("Heading 1", 20, 20, 6, BLACK),
        ("Heading 2", 16, 18, 6, BLACK),
        ("Heading 3", 14, 16, 4, "434343"),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(9.5)
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.right_indent = Inches(0.25)
    code.paragraph_format.space_before = Pt(6)
    code.paragraph_format.space_after = Pt(10)
    code.paragraph_format.line_spacing = 1.0

    caption = styles["Caption"]
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    caption.font.size = Pt(9.5)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False

    table_text = styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    table_text.font.name = "Arial"
    table_text._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    table_text._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    table_text.font.size = Pt(9.2)
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.08

    table_header = styles.add_style("Table Header", WD_STYLE_TYPE.PARAGRAPH)
    table_header.font.name = "Arial"
    table_header._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    table_header._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    table_header.font.size = Pt(9.2)
    table_header.font.bold = True
    table_header.paragraph_format.space_before = Pt(0)
    table_header.paragraph_format.space_after = Pt(0)
    table_header.paragraph_format.line_spacing = 1.08


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.extend([r_fonts, color, underline])
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\[[^\]]+\]\([^)]+\)|\*\*.+?\*\*|`.+?`|https?://\S+)")


def add_inline(paragraph, text: str) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("["):
            label, url = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        elif token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`"):
            set_font(paragraph.add_run(token[1:-1]), name="Consolas", size=9.5, color="333333")
        else:
            url = token.rstrip(".,)")
            suffix = token[len(url):]
            add_hyperlink(paragraph, url, url)
            if suffix:
                paragraph.add_run(suffix)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_shading(paragraph, fill=LIGHT) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if cols == 4:
        widths = [3600, 1500, 2700, 1560]
    elif cols == 3:
        widths = [2100, 3300, 3960]
    elif cols == 2:
        widths = [2700, 6660]
    else:
        widths = [9360 // cols] * cols
        widths[-1] += 9360 - sum(widths)
    numeric_result_table = rows[0][0].startswith("Preregistered")
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.style = "Table Header" if r_idx == 0 else "Table Text"
            if numeric_result_table and c_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, value)
        if r_idx == 0:
            set_repeat_table_header(table.rows[r_idx])
    set_table_geometry(table, widths, indent_dxa=0)
    set_table_borders(table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(3)


def load_figures() -> dict[str, Path]:
    figures = {
        "behavioral": FIGURES_DIR / "behavioral-results.png",
        "logprob": FIGURES_DIR / "logprob-results.png",
    }
    for label, path in figures.items():
        if not path.is_file():
            raise FileNotFoundError(f"Missing {label} figure: {path}")
        with Image.open(path) as image:
            image.verify()
    return figures


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        parts = [p.strip() for p in lines[idx].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-+:?", p) for p in parts):
            rows.append(parts)
        idx += 1
    return rows, idx


def build_doc(output_path: Path = DEFAULT_OUTPUT) -> Path:
    figures = load_figures()
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    idx = 0
    in_code = False
    code_lines: list[str] = []
    seen_title = False
    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = doc.add_paragraph(style="Code Block")
                add_shading(p)
                run = p.add_run("\n".join(code_lines))
                set_font(run, name="Consolas", size=9.5, color="222222")
                in_code = False
            idx += 1
            continue
        if in_code:
            code_lines.append(raw)
            idx += 1
            continue
        if not stripped:
            idx += 1
            continue
        if stripped == "{{FIGURE:behavioral}}" or stripped.startswith("![Behavioral alignment results:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            shape = p.add_run().add_picture(str(figures["behavioral"]), width=Inches(6.45))
            shape._inline.docPr.set("descr", "Bar chart: all strictly valid exact source user-message and mirrored-conflict outputs follow the user's requested parity; only the explicit maximize control follows displayed reward.")
            shape._inline.docPr.set("title", "Behavioral alignment results")
            idx += 1
            continue
        if stripped == "{{FIGURE:logprob}}" or stripped.startswith("![Direct-probability diagnostics:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            shape = p.add_run().add_picture(str(figures["logprob"]), width=Inches(6.45))
            shape._inline.docPr.set("descr", "Forest plot: the reward-aligned controllability interaction is negative, the current-minus-archived contrast is negative, and the transform reversal is near zero; positive values were predicted by reward optimization.")
            shape._inline.docPr.set("title", "Direct-probability diagnostics")
            idx += 1
            continue
        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue
        if stripped.startswith("# ") and not seen_title:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[2:])
            set_font(run, size=26, bold=False)
            seen_title = True
            idx += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:], style="Heading 1")
            idx += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(stripped[4:], style="Heading 2")
            idx += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            add_inline(p, text)
            idx += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            idx += 1
            continue
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[1:-1])
            set_font(run, size=14, italic=True, color="434343")
            idx += 1
            continue
        p = doc.add_paragraph()
        add_inline(p, stripped)
        if stripped.startswith("Prepared for") or re.fullmatch(r"\d{1,2} \w+ 20\d{2}", stripped):
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            for run in p.runs:
                set_font(run, size=10, color=MUTED)
        if stripped.startswith("**Figure"):
            p.style = doc.styles["Caption"]
            p.paragraph_format.keep_with_next = False
            for run in p.runs:
                set_font(run, size=9.5, color=MUTED)
        idx += 1

    doc.core_properties.title = "When Does a Grader Hint Become an Objective?"
    doc.core_properties.subject = "SPAR Model Forensics take-home: Odd Number environment"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "model forensics, reward hacking, odd number environment, Qwen3"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a non-authoritative DOCX from takehome.md.")
    parser.add_argument(
        "--output",
        type=Path,
        default=DEFAULT_OUTPUT,
        help=f"Output path (default: {DEFAULT_OUTPUT})",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Allow replacement of an existing output, including the checked-in audited DOCX.",
    )
    parser.add_argument(
        "--check-output-policy",
        action="store_true",
        help="Validate output selection without loading build dependencies or writing files.",
    )
    return parser.parse_args()


def prepare_output_path(output: Path, force: bool) -> Path:
    resolved = output.resolve()
    if resolved == AUDITED_DOCX.resolve() and not force:
        raise ValueError("Refusing to replace the checked-in audited DOCX without --force.")
    if resolved.exists() and not force:
        raise ValueError(f"Refusing to replace existing output without --force: {resolved}")
    return resolved


def main() -> None:
    args = parse_args()
    try:
        output = prepare_output_path(args.output, args.force)
    except ValueError as exc:
        raise SystemExit(str(exc)) from exc
    if args.check_output_policy:
        print(f"Output policy OK: {output}")
        return
    load_build_dependencies()
    print(build_doc(output))


if __name__ == "__main__":
    main()
